"""
Servicio para procesamiento de documentos PDF y DOCX
Extracción de texto y fragmentación (chunking)
"""
import logging
from typing import List, Dict, Tuple
from pathlib import Path
import uuid

# Procesamiento de PDF
import fitz  # PyMuPDF

# Procesamiento de DOCX
from docx import Document as DocxDocument

from app.core.config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Procesador de documentos PDF y DOCX"""

    def __init__(self):
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
        logger.info(f"DocumentProcessor inicializado - Chunk size: {self.chunk_size}, Overlap: {self.chunk_overlap}")

    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extraer texto de un archivo PDF

        Args:
            file_path: Ruta al archivo PDF

        Returns:
            Texto extraído del PDF
        """
        try:
            doc = fitz.open(file_path)
            text = ""

            for page_num in range(len(doc)):
                page = doc[page_num]
                text += page.get_text()

            doc.close()

            logger.info(f"Texto extraído del PDF: {len(text)} caracteres")
            return text

        except Exception as e:
            logger.error(f"Error al extraer texto de PDF {file_path}: {e}")
            raise Exception(f"Error al procesar PDF: {str(e)}")

    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extraer texto de un archivo DOCX

        Args:
            file_path: Ruta al archivo DOCX

        Returns:
            Texto extraído del DOCX
        """
        try:
            doc = DocxDocument(file_path)
            text = ""

            # Extraer texto de párrafos
            for para in doc.paragraphs:
                text += para.text + "\n"

            # Extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"

            logger.info(f"Texto extraído del DOCX: {len(text)} caracteres")
            return text

        except Exception as e:
            logger.error(f"Error al extraer texto de DOCX {file_path}: {e}")
            raise Exception(f"Error al procesar DOCX: {str(e)}")

    def extract_text(self, file_path: str, file_extension: str) -> str:
        """
        Extraer texto de un documento (PDF o DOCX)

        Args:
            file_path: Ruta al archivo
            file_extension: Extensión del archivo ('pdf' o 'docx')

        Returns:
            Texto extraído
        """
        if file_extension.lower() == 'pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension.lower() in ['docx', 'doc']:
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Extensión no soportada: {file_extension}")

    def chunk_text(self, text: str) -> List[str]:
        """
        Dividir texto en fragmentos (chunks) con solapamiento

        Args:
            text: Texto completo a fragmentar

        Returns:
            Lista de fragmentos de texto
        """
        # Limpiar texto
        text = text.strip()

        # Dividir en palabras
        words = text.split()

        chunks = []
        start = 0

        while start < len(words):
            # Tomar chunk_size palabras
            end = start + self.chunk_size
            chunk_words = words[start:end]
            chunk = " ".join(chunk_words)

            chunks.append(chunk)

            # Avanzar con solapamiento
            start += self.chunk_size - self.chunk_overlap

            # Evitar bucle infinito si chunk_size < chunk_overlap
            if start >= len(words):
                break

        logger.info(f"Texto fragmentado en {len(chunks)} chunks")
        return chunks

    def process_document(
        self,
        file_path: str,
        filename: str,
        file_extension: str,
        document_id: str,
        category: str
    ) -> Tuple[List[str], List[Dict], List[str]]:
        """
        Procesar documento completo: extraer texto y fragmentar

        Args:
            file_path: Ruta al archivo
            filename: Nombre del archivo
            file_extension: Extensión del archivo
            document_id: ID del documento en la BD
            category: Categoría del documento

        Returns:
            Tupla (chunks, metadatas, ids)
            - chunks: Lista de fragmentos de texto
            - metadatas: Lista de metadatos para cada fragmento
            - ids: Lista de IDs únicos para cada fragmento
        """
        # 1. Extraer texto
        text = self.extract_text(file_path, file_extension)

        if not text or len(text.strip()) < 10:
            raise ValueError("El documento está vacío o no contiene texto extraíble")

        # 2. Fragmentar texto
        chunks = self.chunk_text(text)

        # 3. Crear metadatos e IDs para cada chunk
        metadatas = []
        chunk_ids = []

        for i, chunk in enumerate(chunks):
            chunk_id = str(uuid.uuid4())
            metadata = {
                "doc_id": document_id,
                "filename": filename,
                "category": category,
                "chunk_index": i,
                "chunk_total": len(chunks)
            }

            chunk_ids.append(chunk_id)
            metadatas.append(metadata)

        logger.info(f"Documento procesado: {len(chunks)} chunks generados")

        return chunks, metadatas, chunk_ids

    def validate_file_size(self, file_size_bytes: int) -> bool:
        """
        Validar tamaño del archivo

        Args:
            file_size_bytes: Tamaño del archivo en bytes

        Returns:
            True si es válido, False si no
        """
        max_size_bytes = settings.MAX_FILE_SIZE_MB * 1024 * 1024

        if file_size_bytes > max_size_bytes:
            logger.warning(f"Archivo demasiado grande: {file_size_bytes / 1024 / 1024:.2f} MB")
            return False

        return True

    def validate_file_extension(self, filename: str) -> bool:
        """
        Validar extensión del archivo

        Args:
            filename: Nombre del archivo

        Returns:
            True si es válida, False si no
        """
        extension = Path(filename).suffix.lstrip('.').lower()
        return extension in [ext.lower() for ext in settings.ALLOWED_EXTENSIONS]


# Instancia global del servicio
document_processor = DocumentProcessor()
